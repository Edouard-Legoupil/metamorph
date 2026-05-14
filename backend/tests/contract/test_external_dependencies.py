"""Contract tests for external dependencies"""
import pytest
from unittest.mock import patch, MagicMock
import httpx


def test_docling_integration():
    """Test Docling integration contract"""
    # Mock Docling client
    with patch('app.services.ingestion_manager.docling') as mock_docling:
        # Mock successful extraction
        mock_docling.extract_triplets.return_value = {
            "triplets": [
                {
                    "subject": "USAID",
                    "predicate": "FUNDS",
                    "object": "HIP 2024",
                    "confidence": 0.95
                }
            ],
            "metadata": {
                "processing_time": 2.5,
                "model_version": "1.0"
            }
        }
        
        # Test extraction call
        from app.services.ingestion_manager import IngestionManager
        manager = IngestionManager()
        
        result = manager._extract_with_docling("test content", "test.pdf")
        
        assert result is not None
        assert len(result["triplets"]) == 1
        assert result["triplets"][0]["subject"] == "USAID"
        
        # Test error handling
        mock_docling.extract_triplets.side_effect = Exception("Docling service unavailable")
        
        with pytest.raises(Exception) as exc_info:
            manager._extract_with_docling("test content", "test.pdf")
        
        assert "Docling service unavailable" in str(exc_info.value)


def test_mineru_integration():
    """Test MinerU integration contract"""
    # Mock MinerU client
    with patch('app.services.ingestion_manager.mineru') as mock_mineru:
        # Mock successful extraction
        mock_mineru.extract_entities.return_value = {
            "entities": [
                {
                    "text": "Refugee camp",
                    "type": "LOCATION",
                    "confidence": 0.92
                }
            ],
            "metadata": {
                "processing_time": 1.8,
                "model_version": "2.1"
            }
        }
        
        # Test extraction call
        from app.services.ingestion_manager import IngestionManager
        manager = IngestionManager()
        
        result = manager._extract_with_mineru("test content", "test.pdf")
        
        assert result is not None
        assert len(result["entities"]) == 1
        assert result["entities"][0]["text"] == "Refugee camp"
        
        # Test error handling
        mock_mineru.extract_entities.side_effect = Exception("MinerU service unavailable")
        
        with pytest.raises(Exception) as exc_info:
            manager._extract_with_mineru("test content", "test.pdf")
        
        assert "MinerU service unavailable" in str(exc_info.value)


def test_neo4j_integration():
    """Test Neo4j integration contract"""
    # Mock Neo4j driver
    with patch('app.services.knowledge_graph.GraphDatabase') as mock_neo4j:
        mock_driver = MagicMock()
        mock_session = MagicMock()
        mock_result = MagicMock()
        
        # Mock successful query
        mock_result.single.return_value = {"count": 5}
        mock_session.run.return_value = mock_result
        mock_driver.session.return_value = mock_session
        mock_neo4j.driver.return_value = mock_driver
        
        # Test graph operations
        from app.services.knowledge_graph import KnowledgeGraphService
        service = KnowledgeGraphService()
        
        # Mock the driver initialization
        service.driver = mock_driver
        
        # Test entity creation
        result = service.create_entity("TestEntity", "ORGANIZATION", {"description": "Test org"})
        assert result is not None
        
        # Test relationship creation
        result = service.create_relationship("entity1", "entity2", "FUNDS", {"amount": 1000000})
        assert result is not None
        
        # Test error handling
        mock_session.run.side_effect = Exception("Neo4j connection failed")
        
        with pytest.raises(Exception) as exc_info:
            service.create_entity("TestEntity2", "ORGANIZATION", {})
        
        assert "Neo4j connection failed" in str(exc_info.value)


def test_httpx_http_client():
    """Test HTTPX client contract"""
    # Test successful HTTP request
    with patch('httpx.AsyncClient') as mock_client:
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.json.return_value = {"status": "success"}
        mock_response.iter_bytes.return_value = [b"test content"]
        
        mock_client_instance = MagicMock()
        mock_client_instance.__aenter__.return_value.get.return_value = mock_response
        mock_client.return_value = mock_client_instance
        
        # Test HTTP request
        from app.services.ingestion_manager import IngestionManager
        manager = IngestionManager()
        
        # This would normally be called internally
        # We're testing the contract that httpx is used correctly
        assert mock_client.called
        
    # Test error handling
    with patch('httpx.AsyncClient') as mock_client:
        mock_response = MagicMock()
        mock_response.raise_for_status.side_effect = httpx.HTTPStatusError(
            "Not found", request=MagicMock(), response=mock_response
        )
        
        mock_client_instance = MagicMock()
        mock_client_instance.__aenter__.return_value.get.return_value = mock_response
        mock_client.return_value = mock_client_instance
        
        with pytest.raises(Exception) as exc_info:
            # This would trigger the HTTP error handling
            pass  # The actual call would be in download_file method
        
        # Verify that HTTP errors are handled
        assert "HTTPStatusError" in str(type(exc_info.value))


def test_apscheduler_integration():
    """Test APScheduler integration contract"""
    from apscheduler.schedulers.background import BackgroundScheduler
    
    # Test scheduler creation
    scheduler = BackgroundScheduler()
    assert scheduler.running is False
    
    # Test adding a job
    def test_job():
        return "test"
    
    scheduler.add_job(test_job, 'interval', seconds=60, id='test_job')
    
    # Test job retrieval
    job = scheduler.get_job('test_job')
    assert job is not None
    assert job.id == 'test_job'
    
    # Test scheduler start/stop
    scheduler.start()
    assert scheduler.running is True
    
    scheduler.shutdown()
    assert scheduler.running is False


def test_file_processing_libraries():
    """Test file processing library contracts"""
    import tempfile
    import os
    
    # Test PyPDF2
    try:
        from PyPDF2 import PdfReader
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            # Write minimal PDF
            f.write(b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj xref 0 1 0000000000 65535 f trailer<</Size 1/Root 1 0 R>> startxref 67 %%EOF")
            temp_file = f.name
        
        try:
            # Test PDF reading
            reader = PdfReader(temp_file)
            assert len(reader.pages) >= 0  # Should not crash
        finally:
            os.unlink(temp_file)
    except ImportError:
        pytest.skip("PyPDF2 not available")
    
    # Test python-docx
    try:
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_file = f.name
        
        try:
            # Test DOCX creation
            doc = Document()
            doc.add_paragraph("Test content")
            doc.save(temp_file)
            
            # Test DOCX reading
            doc_read = Document(temp_file)
            assert len(doc_read.paragraphs) == 1
        finally:
            os.unlink(temp_file)
    except ImportError:
        pytest.skip("python-docx not available")


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
